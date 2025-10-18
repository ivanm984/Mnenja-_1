"""Utilities for creating reporting outputs.

This module now supports both the existing Word compliance report generation
and filling in the Excel *Priloga 10A* template.  The Excel helper replicates
the behaviour described in the product requirements: labels are resolved by
text rather than coordinates, merged cells are handled safely and list values
are rendered as multiline bullet lists.
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, RGBColor
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet


def generate_word_report(
    zahteve: List[Dict[str, Any]],
    results_map: Dict[str, Dict[str, Any]],
    metadata: Dict[str, str],
    output_path: str,
) -> str:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Inches(0.7)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    def resolve_zahteva_label(zahteva: Dict[str, Any]) -> str:
        clen = (zahteva.get("clen") or "").strip()
        naslov = (zahteva.get("naslov") or "").strip()
        if clen and naslov:
            lowered_clen = clen.lower()
            if naslov.lower().startswith(lowered_clen):
                return naslov
            return f"{clen} - {naslov}"
        return naslov or clen or "Neznan pogoj"

    noncompliant_color = RGBColor(0xFF, 0x00, 0x00)

    neskladja = [
        resolve_zahteva_label(zahteva)
        for zahteva in zahteve
        if results_map.get(zahteva["id"], {}).get("skladnost") == "Neskladno"
    ]
    sklepni_status = "NESKLADNA" if neskladja else "SKLADNA"

    doc.add_heading("Sklepna ugotovitev", level=2)
    p = doc.add_paragraph()
    p.add_run("Gradnja po projektu '")
    p.add_run(metadata.get("ime_projekta", "Ni podatka")).italic = True
    p.add_run(
        f"', s številko projekta '{metadata.get('stevilka_projekta', 'Ni podatka')}', "
        f"datumom '{metadata.get('datum_projekta', 'Ni podatka')}' in projektantom "
        f"'{metadata.get('projektant', 'Ni podatka')}', je glede na predloženo dokumentacijo ocenjena kot "
    )
    p.add_run(f"{sklepni_status}").bold = True
    p.add_run(" s prostorskim aktom.")

    if neskladja:
        p = doc.add_paragraph("Ugotovljena so bila neskladja v naslednjih točkah oziroma členih:")
        for tocka in neskladja:
            bullet = doc.add_paragraph(tocka, style="List Bullet")
            for run in bullet.runs:
                run.font.color.rgb = noncompliant_color
    doc.add_paragraph()

    kategorije = {}
    for zahteva in zahteve:
        kategorija = zahteva.get("kategorija", "Ostalo")
        kategorije.setdefault(kategorija, []).append(zahteva)

    preferred_order = [
        "Splošni prostorski izvedbeni pogoji (PIP)",
        "Podrobni prostorski izvedbeni pogoji (PIP NRP)",
        "Podrobni prostorski izvedbeni pogoji (PIP NRP) - Napotilo",
        "Posebni prostorski izvedbeni pogoji (PIP EUP)",
        "Skladnost z Prilogo 1 (Enostavni/Nezahtevni objekti)",
    ]
    final_order = [cat for cat in preferred_order if cat in kategorije]
    final_order.extend([cat for cat in kategorije if cat not in final_order])

    for kategorija in final_order:
        doc.add_heading(kategorija, level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Pogoj"
        hdr_cells[1].text = "Ugotovitve, obrazložitev in dokazila"
        hdr_cells[2].text = "Skladnost in ukrepi"
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        for zahteva in kategorije[kategorija]:
            row_cells = table.add_row().cells
            result = results_map.get(zahteva["id"], {})

            pogoj_p = row_cells[0].paragraphs[0]
            naslov_run = pogoj_p.add_run(resolve_zahteva_label(zahteva))
            naslov_run.bold = True
            pogoj_p.add_run(f"\n\n{zahteva.get('besedilo', 'Brez besedila')}")

            obrazlozitev_p = row_cells[1].paragraphs[0]
            obrazlozitev_p.add_run("Obrazložitev:\n").bold = True
            obrazlozitev_p.add_run(result.get("obrazlozitev", "—"))
            obrazlozitev_p.add_run("\n\nDokazilo v dokumentaciji:\n").bold = True
            obrazlozitev_p.add_run(result.get("evidence", "—"))

            skladnost_p = row_cells[2].paragraphs[0]
            skladnost_p.add_run("Skladnost:\n").bold = True
            skladnost_run = skladnost_p.add_run(result.get("skladnost", "Neznano"))
            ukrep_text = result.get("predlagani_ukrep", "—")
            if ukrep_text and ukrep_text != "—":
                skladnost_p.add_run("\n\nPredlagani ukrepi:\n").bold = True
                skladnost_p.add_run(ukrep_text)

            if result.get("skladnost") == "Neskladno":
                for paragraph in (pogoj_p, obrazlozitev_p):
                    for run in paragraph.runs:
                        run.font.color.rgb = noncompliant_color
                skladnost_run.font.color.rgb = noncompliant_color

    doc.save(output_path)
    return str(Path(output_path).resolve())


# ---------------------------------------------------------------------------
# Excel Priloga 10A helpers
# ---------------------------------------------------------------------------

def _norm(value: Any) -> str:
    """Normalise text for case-insensitive comparisons."""

    if isinstance(value, str):
        return value.strip().lower()
    return ""


def _find_cells(ws: Worksheet, text: str, *, exact: bool = True):
    """Return all cells whose text matches *text*.

    The comparison is case-insensitive.  When ``exact`` is ``False`` the
    function performs a substring search.
    """

    target = _norm(text)
    hits = []
    for row in ws.iter_rows(values_only=False):
        for cell in row:
            if isinstance(cell.value, str):
                value = _norm(cell.value)
                if (value == target) if exact else (target in value):
                    hits.append(cell)
    return hits


def _resolve_top_left_if_merged(ws: Worksheet, row: int, column: int):
    """Return the top-left cell when (row, column) is inside a merge range."""

    for rng in ws.merged_cells.ranges:
        if (rng.min_row <= row <= rng.max_row) and (
            rng.min_col <= column <= rng.max_col
        ):
            return ws.cell(row=rng.min_row, column=rng.min_col)
    return ws.cell(row=row, column=column)


def _write_value_cell_right_of_label(ws: Worksheet, label_cell, value: Any):
    """Write *value* into the value column (column B) for *label_cell*."""

    target = _resolve_top_left_if_merged(ws, label_cell.row, 2)
    target.value = value


def _as_multiline(value: Any) -> str:
    """Convert *value* to multiline text with bullet points for lists."""

    if isinstance(value, list):
        return "\n".join(f"• {str(item)}" for item in value)
    return "" if value is None else str(value)


def fill_priloga10a(
    xlsx_path: str,
    data: Dict[str, Any],
    output_path: Optional[str] = None,
    *,
    sheet_name: str = "10A MNENJE",
) -> str:
    """Fill the Excel *Priloga 10A* template based on the provided ``data``.

    Parameters mirror the standalone helper script that was previously used in
    operations.  The function writes values based on label lookups, supports
    merged cells, boolean flags for ``JE SKLADNA``/``NI SKLADNA`` and renders
    list values as bullet-point multiline text.
    """

    workbook = load_workbook(xlsx_path)
    if sheet_name not in workbook.sheetnames:
        raise ValueError(
            f"List '{sheet_name}' ne obstaja v '{xlsx_path}'. Najdeni listi: {workbook.sheetnames}"
        )
    worksheet = workbook[sheet_name]

    mapping = {
        "naziv mnenja": "naziv_mnenja",
        "navedba mnenjedajalca": "mnenjedajalec_naziv",
        "naslov": "mnenjedajalec_naslov",
        "št. mnenja": "st_mnenja",
        "datum": "datum",
        "predpis oz. podlaga za mnenje": "predpisi",
        "postopek vodil": "postopek_vodil",
        "odgovorna oseba mnenjedajalca": "odgovorna_oseba",
        "naziv gradnje": "naziv_gradnje",
        "kratek opis gradnje": "kratek_opis",
        "številka projekta": "stevilka_projekta",
        "datum izdelave": "datum_izdelave",
        "projektant (naziv družbe)": "projektant",
        "izdelovalec poročila": "izdelovalec_porocila",
        "številka poročila": "stevilka_porocila",
    }

    for label, key in mapping.items():
        hits = _find_cells(worksheet, label, exact=True)
        if not hits:
            continue
        value = _as_multiline(data.get(key, ""))
        _write_value_cell_right_of_label(worksheet, hits[0], value)

    investor_header = _find_cells(worksheet, "INVESTITOR", exact=True)
    if investor_header:
        after_row = investor_header[0].row
        name_cells = [
            cell
            for cell in _find_cells(
                worksheet, "ime in priimek ali naziv družbe", exact=True
            )
            if cell.row > after_row
        ]
        address_cells = [
            cell
            for cell in _find_cells(
                worksheet, "naslov ali poslovni naslov družbe", exact=True
            )
            if cell.row > after_row
        ]

        if len(name_cells) >= 1:
            _write_value_cell_right_of_label(
                worksheet, name_cells[0], data.get("investitor1_ime", "")
            )
        if len(address_cells) >= 1:
            _write_value_cell_right_of_label(
                worksheet, address_cells[0], data.get("investitor1_naslov", "")
            )
        if len(name_cells) >= 2:
            _write_value_cell_right_of_label(
                worksheet, name_cells[1], data.get("investitor2_ime", "")
            )
        if len(address_cells) >= 2:
            _write_value_cell_right_of_label(
                worksheet, address_cells[1], data.get("investitor2_naslov", "")
            )

    pooblascenec_header = _find_cells(worksheet, "POOBLAŠČENEC", exact=True)
    if pooblascenec_header:
        after_row = pooblascenec_header[0].row
        name_cell = next(
            (
                cell
                for cell in _find_cells(
                    worksheet, "ime in priimek ali naziv družbe", exact=True
                )
                if cell.row > after_row
            ),
            None,
        )
        address_cell = next(
            (
                cell
                for cell in _find_cells(
                    worksheet, "naslov ali poslovni naslov družbe", exact=True
                )
                if cell.row > after_row
            ),
            None,
        )
        if name_cell:
            _write_value_cell_right_of_label(
                worksheet, name_cell, data.get("pooblascenec_ime", "")
            )
        if address_cell:
            _write_value_cell_right_of_label(
                worksheet, address_cell, data.get("pooblascenec_naslov", "")
            )

    je_cell = None
    ni_cell = None
    for cell in _find_cells(worksheet, "JE SKLADNA", exact=False):
        je_cell = _resolve_top_left_if_merged(worksheet, cell.row, 2)
    for cell in _find_cells(worksheet, "NI SKLADNA", exact=False):
        ni_cell = _resolve_top_left_if_merged(worksheet, cell.row, 2)

    if "skladna" in data:
        is_ok = bool(data["skladna"])
        if je_cell:
            je_cell.value = is_ok
        if ni_cell:
            ni_cell.value = not is_ok

    for label, key in [
        ("pogoji za PZI", "pogoji_pzi"),
        ("pogoji za izvajanje gradnje", "pogoji_gradnja"),
        ("pogoji za uporabo objekta", "pogoji_uporaba"),
    ]:
        hits = _find_cells(worksheet, label, exact=True)
        if hits:
            _write_value_cell_right_of_label(
                worksheet, hits[0], _as_multiline(data.get(key, ""))
            )

    explanation = _find_cells(
        worksheet,
        "obrazložitev mnenja (strokovna in pravna utemeljitev)",
        exact=True,
    )
    if explanation:
        _write_value_cell_right_of_label(
            worksheet, explanation[0], data.get("obrazlozitev_mnenja", "")
        )

    if not output_path:
        output_path = str(
            Path(xlsx_path).with_name(Path(xlsx_path).stem + "_filled.xlsx")
        )
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)
    return output_path


def _load_data_json(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description="Izpolni Excel prilogo 10A (list '10A MNENJE')."
    )
    parser.add_argument(
        "--input", "-i", required=True, help="Pot do Excel predloge (Priloga10A.xlsx)"
    )
    parser.add_argument(
        "--output",
        "-o",
        required=False,
        help="Pot do izhodne datoteke (privzeto: *_filled.xlsx)",
    )
    parser.add_argument(
        "--sheet",
        default="10A MNENJE",
        help="Ime lista v Excelu (privzeto: '10A MNENJE')",
    )
    parser.add_argument(
        "--data-json",
        "-d",
        help="JSON datoteka s podatki (ključ: vrednost). Če ni podano, uporabi prazen nabor.",
    )
    args = parser.parse_args(argv)

    data: Dict[str, Any] = {}
    if args.data_json:
        data = _load_data_json(args.data_json)

    output = fill_priloga10a(args.input, data, args.output, sheet_name=args.sheet)
    print(output)
    return 0


__all__ = ["generate_word_report", "fill_priloga10a"]


if __name__ == "__main__":
    raise SystemExit(main())
